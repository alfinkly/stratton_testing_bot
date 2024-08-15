from PIL import Image
from pytesseract import pytesseract
from docx import Document
from PIL import Image, ImageGrab, ImageOps
from aiogram.types import InputMediaPhoto, Message
from aiogram.fsm.storage.base import StorageKey
from aiogram.fsm.context import FSMContext
from tg_config import checker_ids
import os
from tg_config import PATHS


def GetDocx(endFilename, data):
    file_name = f"{PATHS['data']}nda.docx"
    document = Document(file_name)

    fio = f"{data.get('Фамилия').capitalize()} {data.get('Имя').capitalize()} {data.get('Отчество').capitalize()}"
    document.paragraphs[
        4].text = f'{fio}, именуемый/ая в дальнейшем как Получающая Сторона, и ТОО “Stratton”, именуемое в дальнейшем (Передающая сторона), в лице Директора Абдиханова Адила Алмасовича, действующего на основании Устава, с другой стороны, заключили настоящее Соглашение о неразглашении конфиденциальной информации (далее – Соглашение) о нижеследующем.'
    document.tables[-1].columns[-1].cells[1].text = fio
    document.tables[-1].columns[-1].cells[3].text = f"ИИН: {data.get('ИИН')} \nEmail:"
    document.save(endFilename)


def Upscale(fileName):
    img_orig = Image.open(fileName)

    img = ImageOps.scale(img_orig, 6, resample=Image.LANCZOS)
    img.save(fileName)


def ReadPhoto(fileName, mainSide=True):
    Upscale(fileName)
    img = Image.open(fileName)
    text = pytesseract.image_to_string(img, lang='rus+kaz').split()

    fields = {'ИИН': None,
              'Имя': None,
              'Фамилия': None,
              'Отчество': None,
              'Дата рождения': None} if mainSide else {'Дата выдачи': None}
    try:
        if mainSide:
            for i in range(len(text)):
                if text[i].lower().find('имя') != -1:
                    fields['Имя'] = text[i + 1].capitalize()
                elif text[i].lower().find('фамилия') != -1:
                    fields['Фамилия'] = text[i + 1].capitalize()
                elif text[i].lower().find('отчество') != -1:
                    fields['Отчество'] = text[i + 1].capitalize()
                elif text[i].lower().find('рождения') != -1:
                    fields['Дата рождения'] = text[i + 1]
                elif text[i].isdigit() and len(text[i]) == 12:
                    fields['ИИН'] = text[i]
        else:
            for i in range(len(text)):
                if text[i].lower().find('действия') != -1:
                    fields['Дата выдачи'] = text[i + 1]
                    break
    except:
        return 0
    finally:
        os.remove(fileName)
    return fields


async def SendToAdmin(msg, state, keyboard):
    data = await state.get_data()

    media = []
    photoPaths = [data['fImageId'], data['sImageId']]
    for path in photoPaths:
        media.append(InputMediaPhoto(media=path))
    await msg.bot.send_media_group(chat_id=7278477437, media=media)

    text1 = f'Пользователь @{msg.from_user.username} отправил данные вам на проверку: '

    text2 = ' \n'.join([f'{k}-{v if v else "нет данных"}' for k, v in data.items() if
                        k not in ['fImageId', 'sImageId', 'face', 'back']])

    storageKey = StorageKey(msg.bot.id, 7278477437, 7278477437)
    opponentState = FSMContext(storage=state.storage, key=storageKey)
    asd = {'text': f'{text1}\n{text2}'}
    asd['fImageId'] = data['fImageId']
    asd['sImageId'] = data['sImageId']
    asd['username'] = msg.from_user.username
    await opponentState.update_data(asd)

    await msg.bot.send_message(7278477437, f'{text1}\n{text2}', reply_markup=keyboard)
    if isinstance(msg, Message):
        await msg.answer(text='Отправлено на проверку админу @dokuzu')
    else:
        await msg.message.answer(text='Отправлено на проверку админу @dokuzu')


async def ReportToAdmin(msg, fields, photoId):
    text = f'Попытка сканирования @{msg.from_user.username} \n'
    if fields:
        for i in fields:
            text += f'{i}-{fields[i] if fields[i] is not None else "нет данных"} \n'
    else:
        text = 'ошибка при сканировании'
    for i in checker_ids:
        await msg.bot.send_photo(i, photo=photoId, disable_notification=True)
        await msg.bot.send_message(i, text, disable_notification=True)
